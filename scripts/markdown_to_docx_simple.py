import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_divider(line: str) -> bool:
    stripped = line.strip().strip("|").replace(" ", "")
    return bool(stripped) and all(ch in "-:" for ch in stripped)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_index, row in enumerate(rows):
        for c_index in range(cols):
            cell = table.cell(r_index, c_index)
            cell.text = row[c_index] if c_index < len(row) else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if r_index == 0:
                        run.bold = True
    document.add_paragraph()


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(mailto:[^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.replace("**", "").replace("*", "").replace("`", "")


def add_paragraph_with_markdown(document: Document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(clean_inline(text))


def convert(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and is_divider(lines[i + 1]):
            rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(document, rows)
            continue

        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            add_paragraph_with_markdown(document, line[2:].strip(), "List Bullet")
        elif line[:3].isdigit() and line[3:5] == ". ":
            add_paragraph_with_markdown(document, line[5:].strip(), "List Number")
        elif len(line) > 3 and line[0].isdigit() and line[1:3] == ". ":
            add_paragraph_with_markdown(document, line[3:].strip(), "List Number")
        else:
            add_paragraph_with_markdown(document, line)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown")
    parser.add_argument("output")
    args = parser.parse_args()
    convert(Path(args.markdown), Path(args.output))


if __name__ == "__main__":
    main()
