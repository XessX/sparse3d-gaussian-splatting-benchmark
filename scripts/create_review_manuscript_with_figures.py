import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from markdown_to_docx_simple import add_table, is_divider, split_table_row, add_paragraph_with_markdown


PROJECT_ROOT = Path(__file__).resolve().parents[1]

FIGURES = [
    (
        "Fig. 1. Benchmark workflow.",
        "Public scenes are split into held-out test views and sparse training views before selection. Degradation experiments alter training images only, while held-out test views remain clean.",
        "fig1_benchmark_workflow.png",
    ),
    (
        "Fig. 2. Clean sparse-view 3DGS performance.",
        "Mean held-out PSNR across view counts and view-selection strategies for the four evaluated scenes.",
        "3dgs_multiscene_psnr_vs_views.png",
    ),
    (
        "Fig. 3. Repeated random-seed variability.",
        "Mean and standard deviation of held-out PSNR across five random seeds, four scenes, and six view counts.",
        "fig3_random_seed_variability.png",
    ),
    (
        "Fig. 4. Five-view selector comparison.",
        "Mean held-out PSNR for quality-only, pose-free image-diverse, pose-aware quality-diverse, random, and uniform 5-view selection.",
        "fig4_selector_comparison_5views.png",
    ),
    (
        "Fig. 5. Degradation sensitivity.",
        "PSNR change relative to clean 5-view baselines for JPEG compression, blur, low-resolution simulation, underexposure, and overexposure.",
        "3dgs_degradation_5views_psnr_drop.png",
    ),
    (
        "Fig. 6. Convergence subset.",
        "Held-out PSNR comparison for 3000 and 7000 3DGS training iterations on selected scenes and splits.",
        "3dgs_convergence_check.png",
    ),
]


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True


def add_markdown(document: Document, markdown_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8-sig").splitlines()
    in_figure_legend_block = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if line == "## Figure Legends":
            in_figure_legend_block = True
            i += 1
            continue
        if in_figure_legend_block and line == "## Table Legends":
            in_figure_legend_block = False
        if in_figure_legend_block:
            i += 1
            continue
        if not line:
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and is_divider(lines[i + 1]):
            rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(document, rows)
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            add_paragraph_with_markdown(document, line[2:].strip(), "List Bullet")
        elif re.match(r"^\d+\.\s", line):
            add_paragraph_with_markdown(document, re.sub(r"^\d+\.\s", "", line), "List Number")
        else:
            add_paragraph_with_markdown(document, line)
        i += 1


def add_figures(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Figures for Review", level=1)
    figure_dir = PROJECT_ROOT / "manuscript" / "figures"
    for title, caption, filename in FIGURES:
        path = figure_dir / filename
        document.add_heading(title, level=2)
        if path.exists():
            document.add_picture(str(path), width=Inches(6.3))
        else:
            document.add_paragraph(f"[Missing figure file: {filename}]")
        document.add_paragraph(caption)


def main() -> None:
    document = Document()
    configure(document)
    add_markdown(document, PROJECT_ROOT / "manuscript" / "submission" / "scientific_reports_manuscript_final.md")
    add_figures(document)
    output = PROJECT_ROOT / "manuscript" / "submission" / "scientific_reports_review_with_figures.docx"
    document.save(output)
    print(f"Saved {output}")


if __name__ == "__main__":
    main()
